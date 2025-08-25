from flask import Blueprint, render_template, request, flash, redirect, url_for, abort, session, jsonify, send_file, current_app
from flask_login import login_required, current_user, current_user
from app.services.employee_service import EmployeeService
from app.models.employee import Employee
from app.models.payroll_entry import PayrollEntry
from app.models.company import Company
from app.models.company_deduction_default import CompanyDeductionDefault
from app.models.employee_recurring_deduction import EmployeeRecurringDeduction
from app.models.beneficiary import Beneficiary
from app.models.employee_medical_aid_info import EmployeeMedicalAidInfo
from app.models.ui19_record import UI19Record
from app.models.document_template import DocumentTemplate
from app.services.payroll_service import (
    calculate_medical_aid_deduction,
    calculate_medical_aid_fringe_benefit,
    calculate_ytd_totals,
)
from app import db
from datetime import datetime, date, timedelta
from decimal import Decimal
from io import BytesIO
import pandas as pd
import re
import hashlib
import tempfile
import os
from docx import Document
from docx2pdf import convert

# Create employees blueprint
employees_bp = Blueprint('employees', __name__, url_prefix='/employees')

# Mapping between spreadsheet headers and model fields
IMPORT_HEADER_MAP = {
    'First Name': 'first_name',
    'Last Name': 'last_name',
    'South African ID Number': 'id_number',
    'Tax Number': 'tax_number',
    'Cell Number': 'cell_number',
    'Email Address': 'email',
    'Date of Birth': 'date_of_birth',
    'Gender': 'gender',
    'Marital Status': 'marital_status',
    'Department': 'department',
    'Job Title': 'job_title',
    'Start Date': 'start_date',
    'Employment Type': 'employment_type',
    'Salary Type': 'salary_type',
    'Hourly Rate': 'hourly_rate',
    'Monthly Salary': 'monthly_salary',
    'Bank Name': 'bank_name',
    'Account Number': 'account_number',
    'Account Type': 'account_type',
    'Annual Leave Days': 'annual_leave_days',
}

# Normalization lookup tables
GENDER_MAP = {
    'm': 'Male',
    'male': 'Male',
    'f': 'Female',
    'female': 'Female'
}

STATUS_MAP = {
    'single': 'Single',
    'married': 'Married',
    'divorced': 'Divorced',
    'widowed': 'Widowed',
}

BANK_MAP = {
    'absa': 'ABSA',
    'fnb': 'FNB',
    'standard bank': 'Standard Bank',
    'capitec': 'Capitec Bank',
    'nedbank': 'Nedbank',
}


def validate_sa_cell_number(cell_number: str) -> tuple[bool, str]:
    """Validate South African cell number format and return validity and error message"""
    if not cell_number:
        return False, "Cell number is required"
    
    # Check if matches the full regex pattern for both formats
    pattern = r'^(\+27|0)[6-8][0-9]{8}$'
    if re.match(pattern, cell_number.strip()):
        return True, ""
    
    return False, "Please enter a valid South African mobile number (0791234567 or +27791234567)."

def normalize_phone(val: str) -> str:
    """Normalize a South African phone number to +27 format"""
    if not val:
        return ''
    
    # Remove all non-digit characters
    digits = re.sub(r'\D', '', str(val))
    
    # Handle different input formats
    if digits.startswith('27') and len(digits) == 11:
        # Already in 27xxxxxxxxx format
        pass
    elif digits.startswith('0') and len(digits) == 10:
        # SA format 0xxxxxxxxx -> convert to 27xxxxxxxxx
        digits = '27' + digits[1:]
    elif len(digits) == 9:
        # Missing leading 0, assume SA format -> add 27
        digits = '27' + digits
    else:
        # Invalid format, return as-is
        return val
    
    return '+' + digits if digits else ''


def parse_date(val):
    """Parse a date from various common formats"""
    if pd.isna(val) or val == '':
        return None
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d'):
        try:
            return datetime.strptime(str(val), fmt).date()
        except Exception:
            continue
    try:
        return pd.to_datetime(val, dayfirst=True).date()
    except Exception:
        return None


def clean_import_value(val):
    """Clean up raw values imported from Excel."""
    if pd.isna(val) or val is None:
        return ''
    val = str(val).strip()
    if val.startswith("'"):
        val = val.lstrip("'")
    return val

def slugify_company_name(company_name):
    """Convert company name to a slug format"""
    # Remove special characters and convert to lowercase
    slug = re.sub(r'[^a-zA-Z0-9\s]', '', company_name.lower())
    # Replace spaces with hyphens
    slug = re.sub(r'\s+', '-', slug.strip())
    return slug

def generate_unique_company_prefix(company_id, company_name):
    """Generate a unique 5-character prefix for the company"""
    # First try to generate from company name
    slug = slugify_company_name(company_name)
    
    # Extract characters for prefix (skip hyphens, take consonants first, then vowels)
    chars = []
    consonants = []
    vowels = []
    
    for char in slug:
        if char.isalpha():
            if char in 'aeiou':
                vowels.append(char.upper())
            else:
                consonants.append(char.upper())
    
    # Prefer consonants first, then vowels
    chars = consonants + vowels
    
    # Try different combinations to get 5 unique characters
    if len(chars) >= 5:
        prefix = ''.join(chars[:5])
    elif len(chars) >= 3:
        # Pad with first letters repeated or use company ID hash
        prefix = ''.join(chars[:3])
        # Add digits from company ID
        prefix += str(company_id).zfill(2)[:2]
    else:
        # Fallback: use hash of company name + ID
        hash_input = f"{company_name}_{company_id}".encode()
        hash_hex = hashlib.md5(hash_input).hexdigest()
        prefix = ''.join(c.upper() for c in hash_hex if c.isalpha())[:5]
        if len(prefix) < 5:
            prefix = prefix + hash_hex[:5-len(prefix)].upper()
    
    # Ensure exactly 5 characters
    prefix = prefix[:5].ljust(5, 'X')
    
    return prefix

def check_prefix_uniqueness(prefix, company_id):
    """Check if prefix is unique across all companies except current one"""
    # Check if any other company is already using this prefix
    existing = db.session.query(Employee.employee_id).filter(
        Employee.employee_id.like(f"{prefix}-EMP%"),
        Employee.company_id != company_id
    ).first()
    
    return existing is None

def validate_south_african_id(id_number):
    """Validate South African ID number format and basic structure"""
    # Check if ID is exactly 13 digits
    if not id_number or len(id_number) != 13 or not id_number.isdigit():
        return False, "ID number must be exactly 13 digits"
    
    # Extract date components (YYMMDD)
    try:
        year = int(id_number[:2])
        month = int(id_number[2:4])
        day = int(id_number[4:6])
        
        # Determine century (if year > 21, assume 19xx, else 20xx)
        full_year = 1900 + year if year > 21 else 2000 + year
        
        # Validate date components - basic range checks
        if month < 1 or month > 12:
            return False, "Invalid month in ID number"
        
        if day < 1 or day > 31:
            return False, "Invalid day in ID number"
        
        # Check if date is not in the future (with some flexibility)
        try:
            test_date = date(full_year, month, day)
            if test_date > date.today():
                return False, "Date of birth cannot be in the future"
        except ValueError:
            # Allow some flexibility for dates that might be edge cases
            pass
            
    except ValueError:
        return False, "Invalid date format in ID number"
    
    # Extract gender sequence (digits 7-10) - basic range check
    try:
        gender_sequence = int(id_number[6:10])
        if gender_sequence < 0 or gender_sequence > 9999:
            return False, "Invalid gender sequence in ID number"
    except ValueError:
        return False, "Invalid gender sequence in ID number"
    
    # Skip checksum validation as it's often problematic with real IDs
    # Focus on format and basic structure validation instead
    
    return True, "Valid South African ID number"

def extract_info_from_id(id_number):
    """Extract date of birth and gender from valid South African ID number"""
    if len(id_number) != 13 or not id_number.isdigit():
        return None, None
    
    # Extract date components
    year = int(id_number[:2])
    month = int(id_number[2:4])
    day = int(id_number[4:6])
    
    # Determine century
    full_year = 1900 + year if year > 21 else 2000 + year
    
    try:
        birth_date = date(full_year, month, day)
    except ValueError:
        return None, None
    
    # Extract gender (digits 7-10: >= 5000 = Male, < 5000 = Female)
    gender_sequence = int(id_number[6:10])
    gender = "Male" if gender_sequence >= 5000 else "Female"
    
    return birth_date, gender

def generate_employee_id(company_id):
    """Generate a unique employee ID for the given company"""
    # Get company details
    company = Company.query.get(company_id)
    if not company:
        raise ValueError("Company not found")
    
    # Generate base prefix
    base_prefix = generate_unique_company_prefix(company_id, company.name)
    
    # Ensure prefix uniqueness across companies
    prefix = base_prefix
    counter = 1
    while not check_prefix_uniqueness(prefix, company_id):
        # If not unique, modify slightly
        if counter == 1:
            prefix = base_prefix[:-1] + str(company_id % 10)
        else:
            # Use hash fallback
            hash_input = f"{company.name}_{company_id}_{counter}".encode()
            hash_hex = hashlib.md5(hash_input).hexdigest()
            prefix = ''.join(c.upper() for c in hash_hex if c.isalpha())[:5]
            if len(prefix) < 5:
                prefix = prefix + hash_hex[:5-len(prefix)].upper()
        counter += 1
        if counter > 10:  # Safety break
            break
    
    # Get next sequence number for this company
    employee_count = Employee.query.filter_by(company_id=company_id).count()
    seq = employee_count + 1
    
    # Generate employee ID and ensure it's unique
    max_attempts = 100
    for attempt in range(max_attempts):
        employee_id = f"{prefix}-EMP{seq:03d}"
        
        # Check if this ID already exists anywhere in the system
        existing = Employee.query.filter_by(employee_id=employee_id).first()
        if not existing:
            return employee_id
        
        seq += 1
    
    # Fallback if all attempts failed
    raise ValueError("Unable to generate unique employee ID after maximum attempts")

@employees_bp.route('/')
@login_required
def index():
    """Employee list page with search and filter capabilities"""
    
    # Get selected company from session
    selected_company_id = session.get('selected_company_id')
    
    # Redirect to company selection if no company is selected
    if not selected_company_id:
        flash('Please select a company to view employees.', 'warning')
        if current_user.is_accountant:
            return redirect(url_for('accountant_dashboard.dashboard'))
        else:
            return redirect(url_for('dashboard.overview'))
    
    # Get query parameters
    search = request.args.get('search', '')
    department = request.args.get('department', '')
    status = request.args.get('status', '')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    # Get filtered employees scoped to selected company
    employees_data = EmployeeService.get_employees_paginated(
        search=search,
        department=department,
        status=status,
        page=page,
        per_page=per_page,
        company_id=selected_company_id
    )
    
    # Get filter options scoped to selected company
    departments = EmployeeService.get_departments(company_id=selected_company_id)
    statuses = EmployeeService.get_employment_statuses()
    
    return render_template('employees/index.html',
                         employees=employees_data['employees'],
                         pagination=employees_data['pagination'],
                         departments=departments,
                         statuses=statuses,
                         current_search=search,
                         current_department=department,
                         current_status=status)

@employees_bp.route('/<int:employee_id>')
@login_required
def view(employee_id):
    """Employee detail view"""
    
    # Get selected company from session
    selected_company_id = session.get('selected_company_id')
    
    employee = EmployeeService.get_employee_by_id(employee_id)
    
    if not employee:
        abort(404)
    
    # Verify employee belongs to selected company
    if selected_company_id and employee.company_id != selected_company_id:
        abort(403)  # Forbidden - employee doesn't belong to current company

    # Calculate time with company
    time_with_company = "Not available"
    if employee.start_date:
        today = date.today()
        delta = today - employee.start_date
        years = delta.days // 365
        months = (delta.days % 365) // 30
        
        if years > 0:
            time_with_company = f"{years} year{'s' if years != 1 else ''}"
            if months > 0:
                time_with_company += f", {months} month{'s' if months != 1 else ''}"
        elif months > 0:
            time_with_company = f"{months} month{'s' if months != 1 else ''}"
        else:
            time_with_company = "Less than 1 month"

    # Load medical aid info record via relationship
    medical_aid_info = employee.medical_aid_info

    # Flag when medical aid deduction is calculated but info missing
    medical_aid_config_missing = False
    if not medical_aid_info:
        calc_deduction = (
            EmployeeRecurringDeduction.query
            .join(Beneficiary)
            .filter(
                EmployeeRecurringDeduction.employee_id == employee.id,
                EmployeeRecurringDeduction.is_active.is_(True),
                EmployeeRecurringDeduction.amount_type == 'Calculated',
                Beneficiary.type == 'Medical Aid'
            )
            .first()
        )
        medical_aid_config_missing = calc_deduction is not None

    # medical_aid_deduction_conflict_check
    medical_aid_conflict = False
    if employee.recurring_deductions:
        active_medical_deductions = (
            EmployeeRecurringDeduction.query
            .join(Beneficiary)
            .filter(
                EmployeeRecurringDeduction.employee_id == employee.id,
                EmployeeRecurringDeduction.is_active.is_(True),
                Beneficiary.type == 'Medical Aid',
            )
            .count()
        )
        medical_aid_conflict = active_medical_deductions > 1

    medical_beneficiaries = []
    if selected_company_id:
        medical_beneficiaries = Beneficiary.query.filter_by(
            company_id=selected_company_id,
            type='Medical Aid'
        ).order_by(Beneficiary.name.asc()).all()

    default_medical_beneficiary_id = employee.linked_medical_beneficiary_id
    if not default_medical_beneficiary_id:
        deduction = (
            EmployeeRecurringDeduction.query
            .join(Beneficiary)
            .filter(
                EmployeeRecurringDeduction.employee_id == employee.id,
                EmployeeRecurringDeduction.is_active.is_(True),
                Beneficiary.type == 'Medical Aid',
            )
            .order_by(EmployeeRecurringDeduction.id)
            .first()
        )
        if deduction:
            default_medical_beneficiary_id = deduction.beneficiary_id

    # Calculate medical aid deduction once for this view
    medical_aid_deduction_amount = calculate_medical_aid_deduction(employee)

    employee_medical_aid_info = medical_aid_info

    # Load UI19 termination records for this employee
    ui19_records = UI19Record.query.filter_by(employee_id=employee.id).order_by(UI19Record.created_at.desc()).all()
    
    return render_template(
        'employees/view.html',
        employee=employee,
        time_with_company=time_with_company,
        medical_beneficiaries=medical_beneficiaries,
        default_medical_beneficiary_id=default_medical_beneficiary_id,
        medical_aid_conflict=medical_aid_conflict,
        medical_aid_deduction_amount=medical_aid_deduction_amount,
        employee_medical_aid_info=employee_medical_aid_info,
        medical_aid_config_missing=medical_aid_config_missing,
        ui19_records=ui19_records,
        today=date.today()
    )


@employees_bp.route('/search')
@login_required
def search():
    """AJAX endpoint for employee search"""
    
    # Get selected company from session
    selected_company_id = session.get('selected_company_id')
    
    query = request.args.get('q', '')
    
    if len(query) < 2:
        return {'employees': []}
    
    employees = EmployeeService.search_employees(query, limit=10, company_id=selected_company_id)
    
    return {
        'employees': [emp.to_dict() for emp in employees]
    }


@employees_bp.route('/download-import-template')
@login_required
def download_import_template():
    """Download employee import template as XLSX"""
    # Column headers for the import template
    columns = [
        "First Name", "Last Name", "South African ID Number", "Tax Number", "Cell Number", "Email Address",
        "Date of Birth", "Gender", "Marital Status", "Department", "Job Title", "Start Date",
        "Employment Status", "Salary Type", "Hourly Rate", "Monthly Salary",
        "Bank Name", "Account Number", "Account Type", "Annual Leave Days"
    ]

    # DataFrame with one empty row so Excel shows a blank line under the headers
    df = pd.DataFrame([["" for _ in columns]], columns=columns)

    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Employees")

        workbook = writer.book
        worksheet = writer.sheets["Employees"]
        text_format = workbook.add_format({"num_format": "@"})
        text_columns = [
            "South African ID Number",
            "Tax Number",
            "Cell Number",
            "Account Number",
        ]
        for idx, col in enumerate(columns):
            if col in text_columns:
                worksheet.set_column(idx, idx, None, text_format)

    output.seek(0)


    return send_file(
        BytesIO(output.getvalue()),
        as_attachment=True,
        download_name="employee_import_template.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

@employees_bp.route('/new', methods=['GET', 'POST'])
@login_required
def new():
    """Add new employee"""
    
    # Get selected company from session
    selected_company_id = session.get('selected_company_id')
    
    # Redirect to company selection if no company is selected
    if not selected_company_id:
        flash('Please select a company to add employees.', 'warning')
        if current_user.is_accountant:
            return redirect(url_for('accountant_dashboard.dashboard'))
        else:
            return redirect(url_for('dashboard.overview'))
    
    if request.method == 'POST':
        try:
            # Auto-generate unique employee ID
            generated_employee_id = generate_employee_id(selected_company_id)
        except ValueError as e:
            flash(f'Error generating employee ID: {str(e)}', 'error')
            return render_template('employees/new.html')
        
        # Get form data - Personal Information
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        identification_type = request.form.get('identification_type', 'sa_id').strip()
        id_number = request.form.get('id_number', '').strip()
        passport_number = request.form.get('passport_number', '').strip()
        date_of_birth = request.form.get('date_of_birth', '').strip()
        gender = request.form.get('gender', '').strip()
        marital_status = request.form.get('marital_status', '').strip()
        tax_number = request.form.get('tax_number', '').strip()
        cell_number = request.form.get('cell_number', '').strip()
        email = request.form.get('email', '').strip()
        physical_address = request.form.get('physical_address', '').strip()
        
        # Employment Details
        department = request.form.get('department', '').strip()
        custom_department = request.form.get('custom_department', '').strip()
        
        # Handle custom department logic
        if department == 'Other' and custom_department:
            # Create or get the custom department for this company
            from app.models.company_department import CompanyDepartment
            CompanyDepartment.get_or_create_department(selected_company_id, custom_department)
            department = custom_department
        elif department == 'Other' and not custom_department:
            errors.append("Custom department name is required when 'Other' is selected")
        
        job_title = request.form.get('job_title', '').strip()
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        employment_type = request.form.get('employment_type', '').strip()
        employment_status = request.form.get("employment_status", "Active").strip()
        reporting_manager = request.form.get('reporting_manager', '').strip()
        union_member = request.form.get('union_member') == '1'
        union_name = request.form.get('union_name', '').strip()

        
        # Compensation
        salary_type = request.form.get('salary_type', '').strip()
        salary = request.form.get('salary', '').strip()
        overtime_eligible = request.form.get('overtime_eligible') == '1'
        allowances = request.form.get('allowances', '0').strip()
        bonus_type = request.form.get('bonus_type', '').strip()
        
        # Payroll configuration fields
        ordinary_hours_per_day = request.form.get('ordinary_hours_per_day', '').strip()
        work_days_per_month = request.form.get('work_days_per_month', '').strip()
        overtime_calc_method = request.form.get('overtime_calc_method', '').strip()
        overtime_multiplier = request.form.get('overtime_multiplier', '').strip()
        sunday_multiplier = request.form.get('sunday_multiplier', '').strip()
        holiday_multiplier = request.form.get('holiday_multiplier', '').strip()
        
        # Statutory Deductions
        uif_contributing = request.form.get('uif_contributing') == '1'
        sdl_contributing = request.form.get('sdl_contributing') == '1'
        paye_exempt = request.form.get('paye_exempt') == '1'
        
        # Medical Aid
        medical_aid_scheme = request.form.get('medical_aid_scheme', '').strip()
        medical_aid_number = request.form.get('medical_aid_number', '').strip()
        medical_aid_principal_member = request.form.get('medical_aid_principal_member') == '1'

        medical_aid_employer = request.form.get('medical_aid_employer', '').strip()
        medical_aid_employee = request.form.get('medical_aid_employee', '').strip()
        medical_aid_dependants = request.form.get('medical_aid_dependants', '0').strip()
        
        # Leave & Bank
        annual_leave_days = request.form.get('annual_leave_days', '').strip()
        bank_name = request.form.get('bank_name', '').strip()
        account_number = request.form.get('account_number', '').strip()
        account_type = request.form.get('account_type', '').strip()
        
        # Initialize errors list for validation
        errors = []
        if not first_name:
            errors.append("First name is required")
        if not last_name:
            errors.append("Last name is required")
        
        # Validate identification based on type
        if identification_type == 'sa_id':
            if not id_number:
                errors.append("South African ID number is required")
            else:
                # Validate South African ID number format, date, and checksum
                id_valid, id_error_message = validate_south_african_id(id_number)
                if not id_valid:
                    errors.append(id_error_message)
                else:
                    # Auto-extract date of birth and gender from valid ID
                    extracted_birth_date, extracted_gender = extract_info_from_id(id_number)
                    if extracted_birth_date and not date_of_birth:
                        date_of_birth = extracted_birth_date.strftime('%Y-%m-%d')
                    if extracted_gender and not gender:
                        gender = extracted_gender
                # Clear passport number if SA ID is used
                passport_number = ''
        else:  # passport
            if not passport_number:
                errors.append("Passport number is required")
            if not date_of_birth:
                errors.append("Date of birth is required when using passport")
            if not gender:
                errors.append("Gender is required when using passport")
            # Clear SA ID number if passport is used
            id_number = ''
        
        # Validate cell number
        if cell_number:
            cell_valid, cell_error = validate_sa_cell_number(cell_number)
            if not cell_valid:
                errors.append(cell_error)
        else:
            errors.append("Cell number is required")
        if not physical_address:
            errors.append("Physical address is required")
        if not department:
            errors.append("Department is required")
        if not job_title:
            errors.append("Job title is required")
        if not employment_type:
            errors.append("Employment status is required")
        if not salary:
            errors.append("Salary is required")
        if not salary_type or salary_type not in ['monthly', 'hourly', 'daily', 'piece']:
            errors.append("Valid salary type is required")
        if not start_date:
            errors.append("Start date is required")
        if not annual_leave_days:
            errors.append("Annual leave days is required")
        if not bank_name:
            errors.append("Bank name is required")
        if not account_number:
            errors.append("Account number is required")
        if not account_type:
            errors.append("Account type is required")
        
        # Validate unique fields within company scope
        if id_number and Employee.query.filter_by(company_id=selected_company_id, id_number=id_number).first():
            errors.append("ID number already exists in this company")
        if tax_number and Employee.query.filter_by(tax_number=tax_number).first():
            errors.append("Tax number already exists")
        
        # Validate numeric salary
        salary_decimal = None
        try:
            salary_decimal = Decimal(salary)
            if salary_decimal <= 0:
                errors.append("Salary must be greater than 0")
        except:
            errors.append("Invalid salary amount")
        
        # Validate dates and numeric fields
        start_date_obj = None
        end_date_obj = None
        date_of_birth_obj = None
        
        try:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
        except:
            errors.append("Invalid start date format")
            
        if end_date:
            try:
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            except:
                errors.append("Invalid end date format")
                
        if date_of_birth:
            try:
                date_of_birth_obj = datetime.strptime(date_of_birth, '%Y-%m-%d').date()
            except:
                errors.append("Invalid date of birth format")
                
        # Validate numeric fields
        allowances_decimal = Decimal('0')
        if allowances:
            try:
                allowances_decimal = Decimal(allowances)
                if allowances_decimal < 0:
                    errors.append("Allowances cannot be negative")
            except:
                errors.append("Invalid allowances amount")
                
        annual_leave_days_int = 15
        if annual_leave_days:
            try:
                annual_leave_days_int = int(annual_leave_days)
                if annual_leave_days_int < 0 or annual_leave_days_int > 30:
                    errors.append("Annual leave days must be between 0 and 30")
            except:
                errors.append("Invalid annual leave days")
                
        medical_aid_dependants_int = 0
        if medical_aid_dependants:
            try:
                medical_aid_dependants_int = int(medical_aid_dependants)
                if medical_aid_dependants_int < 0:
                    errors.append("Medical aid dependants cannot be negative")
            except:
                errors.append("Invalid medical aid dependants")
                
        # Get company for defaults
        company = Company.query.get(selected_company_id)
        
        # Validate payroll configuration fields with company defaults
        ordinary_hours_per_day_decimal = None
        if ordinary_hours_per_day:
            try:
                ordinary_hours_per_day_decimal = Decimal(ordinary_hours_per_day)
                if ordinary_hours_per_day_decimal <= 0 or ordinary_hours_per_day_decimal > 24:
                    errors.append("Ordinary hours per day must be between 1 and 24")
            except:
                errors.append("Invalid ordinary hours per day")
        elif company and company.default_ordinary_hours_per_day:
            ordinary_hours_per_day_decimal = company.default_ordinary_hours_per_day
        else:
            ordinary_hours_per_day_decimal = Decimal('8')

        work_days_per_month_decimal = None
        if work_days_per_month:
            try:
                work_days_per_month_decimal = Decimal(work_days_per_month)
                if work_days_per_month_decimal <= 0 or work_days_per_month_decimal > 31:
                    errors.append("Work days per month must be between 0.01 and 31")
            except:
                errors.append("Invalid work days per month")
        elif company and company.default_work_days_per_month:
            work_days_per_month_decimal = company.default_work_days_per_month
        else:
            work_days_per_month_decimal = Decimal('22')

        overtime_calc_method_val = overtime_calc_method if overtime_calc_method in ['per_hour', 'per_day', 'fixed_amount'] else 'per_hour'
                
        overtime_multiplier_decimal = None
        if overtime_multiplier:
            try:
                overtime_multiplier_decimal = Decimal(overtime_multiplier)
                if overtime_multiplier_decimal < 1.0 or overtime_multiplier_decimal > 5.0:
                    errors.append("Overtime multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid overtime multiplier")
        elif company and company.overtime_multiplier:
            overtime_multiplier_decimal = company.overtime_multiplier
        else:
            errors.append("Overtime multiplier is required")
                
        sunday_multiplier_decimal = None
        if sunday_multiplier:
            try:
                sunday_multiplier_decimal = Decimal(sunday_multiplier)
                if sunday_multiplier_decimal < 1.0 or sunday_multiplier_decimal > 5.0:
                    errors.append("Sunday multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid sunday multiplier")
        elif company and company.sunday_multiplier:
            sunday_multiplier_decimal = company.sunday_multiplier
        else:
            errors.append("Sunday multiplier is required")
                
        holiday_multiplier_decimal = None
        if holiday_multiplier:
            try:
                holiday_multiplier_decimal = Decimal(holiday_multiplier)
                if holiday_multiplier_decimal < 1.0 or holiday_multiplier_decimal > 5.0:
                    errors.append("Holiday multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid holiday multiplier")
        elif company and company.public_holiday_multiplier:
            holiday_multiplier_decimal = company.public_holiday_multiplier
        else:
            errors.append("Holiday multiplier is required")
                
        medical_aid_employer_decimal = None
        if medical_aid_employer:
            try:
                medical_aid_employer_decimal = Decimal(medical_aid_employer)
                if medical_aid_employer_decimal < 0:
                    errors.append("Medical aid employer contribution cannot be negative")
            except:
                errors.append("Invalid medical aid employer contribution")

        medical_aid_employee_decimal = None
        if medical_aid_employee:
            try:
                medical_aid_employee_decimal = Decimal(medical_aid_employee)
                if medical_aid_employee_decimal < 0:
                    errors.append("Medical aid employee contribution cannot be negative")
            except:
                errors.append("Invalid medical aid employee contribution")
                
        # Auto-calculate fringe benefit as employer contribution (SARS requirement)
        medical_aid_fringe_benefit_decimal = medical_aid_employer_decimal or Decimal('0')
                
        # Validate medical aid number length if provided
        if medical_aid_number and (len(medical_aid_number) < 4 or len(medical_aid_number) > 50):
            errors.append("Medical aid number must be between 4 and 50 characters")
        
        if errors:
            for error in errors:
                flash(error, 'error')
            
            # Preserve recurring deduction data from form submission
            form_deductions = []
            deduction_beneficiary_ids = request.form.getlist('deduction_beneficiary_id[]')
            deduction_amount_types = request.form.getlist('deduction_amount_type[]')
            deduction_values = request.form.getlist('deduction_value[]')
            deductions_populated = request.form.get('deductions_populated', '0') == '1'
            
            # User interacted with deductions if: table was populated with defaults OR form has deduction fields
            customize_deductions_flag = deductions_populated or bool(deduction_beneficiary_ids)
            
            if customize_deductions_flag:
                # Preserve exactly what user submitted (including empty state if all removed)
                for i, beneficiary_id in enumerate(deduction_beneficiary_ids):
                    if beneficiary_id and beneficiary_id.strip():  # Only valid beneficiary IDs
                        form_deductions.append({
                            'beneficiary_id': beneficiary_id,
                            'enabled': True,
                            'amount_type': deduction_amount_types[i] if i < len(deduction_amount_types) else 'fixed',
                            'amount': deduction_values[i] if i < len(deduction_values) else '0.00'
                        })
            
            # Get current company and deduction defaults for form pre-population
            company = Company.query.get(selected_company_id)
            deduction_defaults = CompanyDeductionDefault.get_company_defaults(selected_company_id)
            
            # Get all beneficiaries for the company for dynamic deduction rows
            from app.models.beneficiary import Beneficiary
            beneficiaries_objs = Beneficiary.query.filter_by(company_id=selected_company_id).all()
            beneficiaries = [{'id': b.id, 'name': b.name, 'type': b.type} for b in beneficiaries_objs]
            
            # Get company-specific departments instead of global departments
            from app.models.company_department import CompanyDepartment
            company_departments = CompanyDepartment.get_company_departments(selected_company_id)
            departments = [dept.name for dept in company_departments]
            
            return render_template('employees/new.html',
                                 departments=departments,
                                 generated_employee_id=generated_employee_id,
                                 company=company,
                                 deduction_defaults=deduction_defaults,
                                 beneficiaries=beneficiaries,
                                 form_data=request.form,
                                 form_deductions=form_deductions,
                                 customize_deductions=customize_deductions_flag)
        
        # Create new employee
        try:
            new_employee = Employee()
            new_employee.company_id = selected_company_id  # Assign to selected company
            new_employee.employee_id = generated_employee_id  # Use auto-generated ID
            
            # Personal Information
            new_employee.first_name = first_name
            new_employee.last_name = last_name
            new_employee.identification_type = identification_type
            new_employee.id_number = id_number or None
            new_employee.passport_number = passport_number or None
            new_employee.date_of_birth = date_of_birth_obj
            new_employee.gender = gender or None
            new_employee.marital_status = marital_status or None
            new_employee.tax_number = tax_number or None
            new_employee.cell_number = normalize_phone(cell_number)
            new_employee.email = email or None
            new_employee.physical_address = physical_address
            
            # Employment Details
            new_employee.department = department
            new_employee.job_title = job_title
            new_employee.start_date = start_date_obj
            new_employee.end_date = end_date_obj
            new_employee.employment_type = employment_type
            new_employee.employment_status = employment_status
            new_employee.reporting_manager = reporting_manager or None
            new_employee.union_member = union_member
            new_employee.union_name = union_name or None

            
            # Compensation
            new_employee.salary_type = salary_type
            new_employee.salary = salary_decimal
            new_employee.overtime_eligible = overtime_eligible
            new_employee.allowances = allowances_decimal
            new_employee.bonus_type = bonus_type or None
            
            # Payroll configuration
            new_employee.ordinary_hours_per_day = ordinary_hours_per_day_decimal
            new_employee.work_days_per_month = work_days_per_month_decimal
            new_employee.overtime_calc_method = overtime_calc_method_val
            new_employee.overtime_multiplier = overtime_multiplier_decimal
            new_employee.sunday_multiplier = sunday_multiplier_decimal
            new_employee.holiday_multiplier = holiday_multiplier_decimal
            
            # Statutory Deductions
            new_employee.uif_contributing = uif_contributing
            new_employee.sdl_contributing = sdl_contributing
            new_employee.paye_exempt = paye_exempt
            
            # Medical Aid
            new_employee.medical_aid_scheme = medical_aid_scheme or None
            new_employee.medical_aid_number = medical_aid_number or None
            new_employee.medical_aid_principal_member = medical_aid_principal_member

            new_employee.medical_aid_employer = medical_aid_employer_decimal
            new_employee.medical_aid_employee = medical_aid_employee_decimal
            new_employee.medical_aid_dependants = medical_aid_dependants_int
            new_employee.medical_aid_fringe_benefit = medical_aid_fringe_benefit_decimal
            
            # Leave & Banking
            new_employee.annual_leave_days = annual_leave_days_int
            new_employee.bank_name = bank_name
            new_employee.account_number = account_number
            new_employee.account_type = account_type
            
            db.session.add(new_employee)
            db.session.flush()  # Flush to get the employee ID
            
            # Handle recurring deductions from form
            deduction_beneficiary_ids = request.form.getlist('deduction_beneficiary_id[]')
            deduction_amount_types = request.form.getlist('deduction_amount_type[]')
            deduction_values = request.form.getlist('deduction_value[]')
            deductions_populated = request.form.get('deductions_populated', '0') == '1'
            defaults_applied = 0
            
            # Check if user interacted with deductions table at all
            # User customized if: table was populated with defaults OR form has valid deduction fields
            valid_beneficiary_ids = [bid for bid in deduction_beneficiary_ids if bid and bid.strip()]
            user_customized_deductions = deductions_populated or bool(valid_beneficiary_ids)
            
            current_app.logger.debug(
                "Deduction processing: beneficiary_ids=%s, valid_ids=%s, populated=%s, user_customized=%s",
                deduction_beneficiary_ids,
                valid_beneficiary_ids,
                deductions_populated,
                user_customized_deductions
            )
            
            if user_customized_deductions:
                # User interacted with deductions table - process exactly what they submitted
                if valid_beneficiary_ids:
                    # User has actual deductions to create
                    for i, beneficiary_id in enumerate(deduction_beneficiary_ids):
                        # Skip empty beneficiary IDs (from new rows that weren't filled)
                        if not beneficiary_id or not beneficiary_id.strip():
                            continue

                        try:
                            recurring_deduction = EmployeeRecurringDeduction()
                            recurring_deduction.employee_id = new_employee.id
                            recurring_deduction.beneficiary_id = int(beneficiary_id)
                            amount_type_raw = deduction_amount_types[i] if i < len(deduction_amount_types) else 'Fixed'
                            amount_type_lower = amount_type_raw.lower()
                            if amount_type_lower.startswith('percent'):
                                recurring_deduction.amount_type = 'Percentage'
                            elif amount_type_lower.startswith('calc'):
                                recurring_deduction.amount_type = 'Calculated'
                            else:
                                recurring_deduction.amount_type = 'Fixed'
                            if recurring_deduction.amount_type == 'Calculated':
                                recurring_deduction.value = None
                            else:
                                recurring_deduction.value = float(deduction_values[i]) if i < len(deduction_values) and deduction_values[i] else 0.00
                            recurring_deduction.notes = "Customized during employee creation"
                            recurring_deduction.is_active = True
                            recurring_deduction.effective_date = new_employee.start_date

                            db.session.add(recurring_deduction)
                            defaults_applied += 1

                        except Exception as e:
                            current_app.logger.warning(
                                "Could not create recurring deduction for beneficiary %s: %s",
                                beneficiary_id,
                                str(e),
                            )
                            continue
                else:
                    # User removed all defaults - create no deductions
                    current_app.logger.debug("User removed all default deductions - creating no deductions")
                    defaults_applied = 0
            else:
                # No form interaction - apply company deduction defaults automatically
                deduction_defaults = CompanyDeductionDefault.get_company_defaults(selected_company_id)
                
                for default in deduction_defaults:
                    try:
                        recurring_deduction = EmployeeRecurringDeduction()
                        recurring_deduction.employee_id = new_employee.id
                        recurring_deduction.beneficiary_id = default.beneficiary_id
                        recurring_deduction.amount_type = default.amount_type.title()
                        if recurring_deduction.amount_type == 'Calculated':
                            recurring_deduction.value = None
                        else:
                            recurring_deduction.value = default.amount
                        recurring_deduction.notes = "Auto-applied from company default"
                        recurring_deduction.is_active = True
                        recurring_deduction.effective_date = new_employee.start_date
                        
                        db.session.add(recurring_deduction)
                        defaults_applied += 1
                        
                    except Exception as e:
                        current_app.logger.warning(
                            "Could not apply default deduction %s: %s",
                            default.id,
                            str(e),
                        )
                        continue
            
            db.session.commit()
            current_app.logger.debug(
                "Recurring deductions committed for employee %s", new_employee.id
            )
            
            # Enhanced success message with deduction defaults info
            success_message = f'Employee {first_name} {last_name} has been successfully added!'
            if defaults_applied > 0:
                success_message += f' {defaults_applied} recurring deduction default{"s" if defaults_applied != 1 else ""} automatically applied.'
            
            flash(success_message, 'success')
            return redirect(url_for('employees.index'))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error("Error creating employee: %s", e)
            flash(f'An error occurred while saving the employee: {str(e)}', 'error')
            
            # Get current company and defaults for error case
            company = Company.query.get(selected_company_id)
            deduction_defaults = CompanyDeductionDefault.get_company_defaults(selected_company_id)
            
            # Get all beneficiaries for the company for dynamic deduction rows
            from app.models.beneficiary import Beneficiary
            beneficiaries_objs = Beneficiary.query.filter_by(company_id=selected_company_id).all()
            beneficiaries = [{'id': b.id, 'name': b.name, 'type': b.type} for b in beneficiaries_objs]
            
            # Get company-specific departments instead of global departments
            from app.models.company_department import CompanyDepartment
            company_departments = CompanyDepartment.get_company_departments(selected_company_id)
            departments = [dept.name for dept in company_departments]
            
            return render_template('employees/new.html',
                                 departments=departments,
                                 generated_employee_id=generated_employee_id,
                                 company=company,
                                 deduction_defaults=deduction_defaults,
                                 beneficiaries=beneficiaries,
                                 form_data=request.form)
    
    # GET request - show form with generated employee ID
    try:
        generated_employee_id = generate_employee_id(selected_company_id)
    except ValueError as e:
        flash(f'Error generating employee ID: {str(e)}', 'error')
        generated_employee_id = None
    
    # Get current company and deduction defaults for form pre-population
    company = Company.query.get(selected_company_id)
    deduction_defaults = CompanyDeductionDefault.get_company_defaults(selected_company_id)
    
    # Get all beneficiaries for the company for dynamic deduction rows
    from app.models.beneficiary import Beneficiary
    beneficiaries_objs = Beneficiary.query.filter_by(company_id=selected_company_id).all()
    beneficiaries = [{'id': b.id, 'name': b.name, 'type': b.type} for b in beneficiaries_objs]
    
    # Get company-specific departments instead of global departments
    from app.models.company_department import CompanyDepartment
    company_departments = CompanyDepartment.get_company_departments(selected_company_id)
    departments = [dept.name for dept in company_departments]
    
    return render_template('employees/new.html',
                         departments=departments,
                         generated_employee_id=generated_employee_id,
                         company=company,
                         deduction_defaults=deduction_defaults,
                         beneficiaries=beneficiaries)

@employees_bp.route('/<int:employee_id>/edit', methods=['GET', 'POST'])
def edit(employee_id):
    """Edit employee"""
    
    current_app.logger.debug(
        "Edit Employee - Method: %s, Employee ID: %s", request.method, employee_id
    )
    current_app.logger.debug("Request headers: %s", dict(request.headers))
    current_app.logger.debug("Request form keys: %s", list(request.form.keys()))
    
    employee = Employee.query.get(employee_id)
    if not employee:
        flash('Employee not found.', 'error')
        return redirect(url_for('employees.index'))

    if request.method == 'POST':
        # Detect if this is an AJAX request from modal
        is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
        current_app.logger.debug("Is AJAX request: %s", is_ajax)
        current_app.logger.debug("X-Requested-With header: %s", request.headers.get('X-Requested-With'))

        company = Company.query.get(employee.company_id)
        
        # Get form data - Personal Information
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        identification_type = request.form.get('identification_type', 'sa_id').strip()
        id_number = request.form.get('id_number', '').strip()
        passport_number = request.form.get('passport_number', '').strip()
        tax_number = request.form.get('tax_number', '').strip()
        cell_number = request.form.get('cell_number', '').strip()
        email = request.form.get('email', '').strip()
        date_of_birth = request.form.get('date_of_birth', '').strip()
        gender = request.form.get('gender', '').strip()
        marital_status = request.form.get('marital_status', '').strip()
        physical_address = request.form.get('physical_address', '').strip()
        
        current_app.logger.debug(
            "Form data received - first_name: %s, last_name: %s",
            first_name,
            last_name,
        )
        
        # Employment Information
        department = request.form.get('department', '').strip()
        custom_department = request.form.get('custom_department', '').strip()
        
        # Initialize errors list for edit form validation
        errors = []
        
        # Handle custom department logic
        if department == 'Other' and custom_department:
            # Create or get the custom department for this company
            from app.models.company_department import CompanyDepartment
            CompanyDepartment.get_or_create_department(selected_company_id, custom_department)
            department = custom_department
        elif department == 'Other' and not custom_department:
            errors.append("Custom department name is required when 'Other' is selected")
        
        job_title = request.form.get('job_title', '').strip()
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        employment_type = request.form.get('employment_type', '').strip()
        employment_status = request.form.get("employment_status", "Active").strip()
        reporting_manager = request.form.get('reporting_manager', '').strip()
        union_member = request.form.get('union_member') == '1'
        union_name = request.form.get('union_name', '').strip()
        
        # Compensation
        salary_type = request.form.get('salary_type', '').strip()
        salary = request.form.get('salary', '').strip()
        overtime_eligible = request.form.get('overtime_eligible') == '1'
        allowances = request.form.get('allowances', '').strip()
        bonus_type = request.form.get('bonus_type', '').strip()

        # Payroll configuration fields
        ordinary_hours_per_day = request.form.get('ordinary_hours_per_day', '').strip()
        work_days_per_month = request.form.get('work_days_per_month', '').strip()
        overtime_calc_method = request.form.get('overtime_calc_method', '').strip()
        overtime_multiplier = request.form.get('overtime_multiplier', '').strip()
        sunday_multiplier = request.form.get('sunday_multiplier', '').strip()
        holiday_multiplier = request.form.get('holiday_multiplier', '').strip()
        
        # Statutory Deductions
        uif_contributing = request.form.get('uif_contributing') == '1'
        sdl_contributing = request.form.get('sdl_contributing') == '1'
        paye_exempt = request.form.get('paye_exempt') == '1'
        
        # Medical Aid
        medical_aid_scheme = request.form.get('medical_aid_scheme', '').strip()
        medical_aid_number = request.form.get('medical_aid_number', '').strip()
        medical_aid_principal_member = request.form.get('medical_aid_principal_member') == '1'
        medical_aid_employer = request.form.get('medical_aid_employer', '').strip()
        medical_aid_employee = request.form.get('medical_aid_employee', '').strip()
        medical_aid_dependants = request.form.get('medical_aid_dependants', '0').strip()
        
        # Leave & Banking
        annual_leave_days = request.form.get('annual_leave_days', '').strip()
        bank_name = request.form.get('bank_name', '').strip()
        account_number = request.form.get('account_number', '').strip()
        account_type = request.form.get('account_type', '').strip()
        
        # Validate required fields
        errors = []
        if not first_name:
            errors.append("First name is required")
        if not last_name:
            errors.append("Last name is required")
        
        # Validate identification based on type
        if identification_type == 'sa_id':
            if not id_number:
                errors.append("South African ID number is required")
            else:
                # Validate South African ID if needed
                id_valid, id_error_message = validate_south_african_id(id_number)
                if not id_valid:
                    errors.append(id_error_message)
            # Clear passport number if SA ID is used
            passport_number = ''
        else:  # passport
            if not passport_number:
                errors.append("Passport number is required")
            # Clear SA ID number if passport is used
            id_number = ''
        
        # Validate cell number
        if cell_number:
            cell_valid, cell_error = validate_sa_cell_number(cell_number)
            if not cell_valid:
                errors.append(cell_error)
        else:
            errors.append("Cell number is required")
        if not department:
            errors.append("Department is required")
        if not job_title:
            errors.append("Job title is required")
        if not salary:
            errors.append("Salary is required")
        if not salary_type or salary_type not in ['monthly', 'hourly', 'daily', 'piece']:
            errors.append("Valid salary type is required")
        if not start_date:
            errors.append("Start date is required")
        if not bank_name:
            errors.append("Bank name is required")
        if not account_number:
            errors.append("Account number is required")
        if not physical_address:
            errors.append("Physical address is required")
        
        # Validate unique fields (excluding current employee and scoped to company)
        # Only check ID number if it was changed
        if id_number and id_number != employee.id_number:
            existing_employee = Employee.query.filter(
                Employee.id_number == id_number,
                Employee.company_id == employee.company_id,
                Employee.id != employee_id
            ).first()
            if existing_employee:
                errors.append("ID number already exists in this company")
        
        # Check tax number uniqueness (global check since tax numbers are globally unique)
        if tax_number and tax_number != employee.tax_number:
            existing_employee = Employee.query.filter(
                Employee.tax_number == tax_number,
                Employee.id != employee_id
            ).first()
            if existing_employee:
                errors.append("Tax number already exists")
        
        # Validate numeric salary
        salary_decimal = None
        try:
            salary_decimal = Decimal(salary)
            if salary_decimal <= 0:
                errors.append("Salary must be greater than 0")
        except:
            errors.append("Invalid salary amount")
        
        # Validate date
        start_date_obj = None
        try:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
        except:
            errors.append("Invalid start date format")
        
        # First validation error check
        if errors:
            if is_ajax:
                return jsonify({
                    'success': False,
                    'errors': errors,
                    'form_data': dict(request.form)
                }), 400
            else:
                for error in errors:
                    flash(error, 'error')
                return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Parse numeric fields
        allowances_decimal = Decimal('0')
        if allowances:
            try:
                allowances_decimal = Decimal(allowances)
            except:
                errors.append("Invalid allowances amount")
                
        annual_leave_days_int = 15
        if annual_leave_days:
            try:
                annual_leave_days_int = int(annual_leave_days)
            except:
                errors.append("Invalid annual leave days")
                
        medical_aid_dependants_int = 0
        if medical_aid_dependants:
            try:
                medical_aid_dependants_int = int(medical_aid_dependants)
            except:
                errors.append("Invalid number of dependants")

        # Payroll configuration defaults and validation
        ordinary_hours_per_day_decimal = employee.ordinary_hours_per_day
        if ordinary_hours_per_day:
            try:
                ordinary_hours_per_day_decimal = Decimal(ordinary_hours_per_day)
                if ordinary_hours_per_day_decimal <= 0 or ordinary_hours_per_day_decimal > 24:
                    errors.append("Ordinary hours per day must be between 1 and 24")
            except:
                errors.append("Invalid ordinary hours per day")
        elif company and company.default_ordinary_hours_per_day:
            ordinary_hours_per_day_decimal = company.default_ordinary_hours_per_day

        work_days_per_month_decimal = employee.work_days_per_month
        if work_days_per_month:
            try:
                work_days_per_month_decimal = Decimal(work_days_per_month)
                if work_days_per_month_decimal <= 0 or work_days_per_month_decimal > 31:
                    errors.append("Work days per month must be between 0.01 and 31")
            except:
                errors.append("Invalid work days per month")
        elif company and company.default_work_days_per_month:
            work_days_per_month_decimal = company.default_work_days_per_month
        overtime_calc_method_val = overtime_calc_method if overtime_calc_method in ['per_hour','per_day','fixed_amount'] else employee.overtime_calc_method

        overtime_multiplier_decimal = employee.overtime_multiplier
        if overtime_multiplier:
            try:
                overtime_multiplier_decimal = Decimal(overtime_multiplier)
                if overtime_multiplier_decimal < 1.0 or overtime_multiplier_decimal > 5.0:
                    errors.append("Overtime multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid overtime multiplier")
        elif overtime_multiplier_decimal is None and company and company.overtime_multiplier:
            overtime_multiplier_decimal = company.overtime_multiplier

        sunday_multiplier_decimal = employee.sunday_multiplier
        if sunday_multiplier:
            try:
                sunday_multiplier_decimal = Decimal(sunday_multiplier)
                if sunday_multiplier_decimal < 1.0 or sunday_multiplier_decimal > 5.0:
                    errors.append("Sunday multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid sunday multiplier")
        elif sunday_multiplier_decimal is None and company and company.sunday_multiplier:
            sunday_multiplier_decimal = company.sunday_multiplier

        holiday_multiplier_decimal = employee.holiday_multiplier
        if holiday_multiplier:
            try:
                holiday_multiplier_decimal = Decimal(holiday_multiplier)
                if holiday_multiplier_decimal < 1.0 or holiday_multiplier_decimal > 5.0:
                    errors.append("Holiday multiplier must be between 1.0 and 5.0")
            except:
                errors.append("Invalid holiday multiplier")
        elif holiday_multiplier_decimal is None and company and company.public_holiday_multiplier:
            holiday_multiplier_decimal = company.public_holiday_multiplier


        medical_aid_employer_decimal = None
        if medical_aid_employer:
            try:
                medical_aid_employer_decimal = Decimal(medical_aid_employer)
            except:
                errors.append("Invalid medical aid employer contribution")

        medical_aid_employee_decimal = None
        if medical_aid_employee:
            try:
                medical_aid_employee_decimal = Decimal(medical_aid_employee)
            except:
                errors.append("Invalid medical aid employee contribution")
                
        # Auto-calculate fringe benefit as employer contribution (SARS requirement)
        medical_aid_fringe_benefit_decimal = medical_aid_employer_decimal or Decimal('0')

        # Parse dates
        date_of_birth_obj = None
        if date_of_birth:
            try:
                date_of_birth_obj = datetime.strptime(date_of_birth, '%Y-%m-%d').date()
            except:
                pass  # Optional field
                
        end_date_obj = None
        if end_date:
            try:
                end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            except:
                pass  # Optional field

        # Second validation error check (after parsing)
        if errors:
            if is_ajax:
                return jsonify({
                    'success': False,
                    'errors': errors,
                    'form_data': dict(request.form)
                }), 400
            else:
                for error in errors:
                    flash(error, 'error')
                return redirect(url_for('employees.view', employee_id=employee_id))

        # Update employee
        try:
            # Personal Information
            employee.first_name = first_name
            employee.last_name = last_name
            employee.identification_type = identification_type
            employee.id_number = id_number or None
            employee.passport_number = passport_number or None
            employee.tax_number = tax_number or None
            employee.cell_number = normalize_phone(cell_number)
            employee.email = email or None
            employee.date_of_birth = date_of_birth_obj
            employee.gender = gender or None
            employee.marital_status = marital_status or None
            employee.physical_address = physical_address
            
            # Employment Information
            employee.department = department
            employee.job_title = job_title
            employee.start_date = start_date_obj
            employee.end_date = end_date_obj
            employee.employment_type = employment_type
            employee.employment_status = employment_status
            employee.reporting_manager = reporting_manager or None
            employee.union_member = union_member
            employee.union_name = union_name or None

            
            # Compensation
            employee.salary_type = salary_type
            employee.salary = salary_decimal
            employee.overtime_eligible = overtime_eligible
            employee.allowances = allowances_decimal
            employee.bonus_type = bonus_type or None

            # Payroll configuration
            employee.ordinary_hours_per_day = ordinary_hours_per_day_decimal
            employee.work_days_per_month = work_days_per_month_decimal
            employee.overtime_calc_method = overtime_calc_method_val
            employee.overtime_multiplier = overtime_multiplier_decimal
            employee.sunday_multiplier = sunday_multiplier_decimal
            employee.holiday_multiplier = holiday_multiplier_decimal
            
            # Statutory Deductions
            employee.uif_contributing = uif_contributing
            employee.sdl_contributing = sdl_contributing
            employee.paye_exempt = paye_exempt
            
            # Medical Aid
            employee.medical_aid_scheme = medical_aid_scheme or None
            employee.medical_aid_number = medical_aid_number or None
            employee.medical_aid_principal_member = medical_aid_principal_member

            employee.medical_aid_employer = medical_aid_employer_decimal
            employee.medical_aid_employee = medical_aid_employee_decimal
            employee.medical_aid_dependants = medical_aid_dependants_int
            employee.medical_aid_fringe_benefit = medical_aid_fringe_benefit_decimal
            
            # Leave & Banking
            employee.annual_leave_days = annual_leave_days_int
            employee.bank_name = bank_name
            employee.account_number = account_number
            employee.account_type = account_type
            
            employee.updated_at = datetime.utcnow()
            
            # Handle recurring deductions if customization is enabled
            customize_deductions = request.form.get('customize_deductions') == 'on'
            if not customize_deductions:
                customize_deductions = any(request.form.getlist('deduction_beneficiary_id[]'))
            deductions_added = 0
            deductions_updated = 0
            deductions_removed = 0
            
            if customize_deductions:
                current_app.logger.debug(
                    "Processing recurring deductions for employee %s",
                    employee_id,
                )
                
                # Process customized deductions from form
                deduction_beneficiary_ids = request.form.getlist('deduction_beneficiary_id[]')
                deduction_amount_types = request.form.getlist('deduction_amount_type[]')
                deduction_values = request.form.getlist('deduction_value[]')
                
                current_app.logger.debug(
                    "Form data - beneficiaries: %s", deduction_beneficiary_ids
                )
                current_app.logger.debug(
                    "Form data - types: %s, values: %s",
                    deduction_amount_types,
                    deduction_values,
                )
                
                # Get all existing active deductions for this employee
                existing_deductions = EmployeeRecurringDeduction.query.filter_by(
                    employee_id=employee_id,
                    is_active=True
                ).all()
                
                # Create a map of existing deductions by beneficiary_id
                existing_map = {str(deduction.beneficiary_id): deduction for deduction in existing_deductions}
                processed_beneficiaries = set()
                
                current_app.logger.debug(
                    "Existing deductions: %s",
                    [
                        f"Beneficiary {d.beneficiary_id}: {d.amount_type} {d.value}"
                        for d in existing_deductions
                    ],
                )
                
                # Process incoming deductions
                for i, beneficiary_id in enumerate(deduction_beneficiary_ids):
                    if not beneficiary_id:
                        continue
                        
                    beneficiary_id_str = str(beneficiary_id)
                    processed_beneficiaries.add(beneficiary_id_str)
                    
                    # Get form values for this deduction
                    amount_type_raw = deduction_amount_types[i] if i < len(deduction_amount_types) else 'Fixed'
                    amount_type_lower = amount_type_raw.lower()
                    if amount_type_lower.startswith('percent'):
                        amount_type = 'Percentage'
                    elif amount_type_lower.startswith('calc'):
                        amount_type = 'Calculated'
                    else:
                        amount_type = 'Fixed'

                    # Properly handle calculated deductions with no numeric value
                    if amount_type == 'Calculated':
                        value = None
                    else:
                        try:
                            value = float(deduction_values[i]) if i < len(deduction_values) and deduction_values[i] else 0.00
                        except (ValueError, IndexError):
                            value = 0.00

                    current_app.logger.debug(
                        "Processing beneficiary %s: type=%s, value=%s",
                        beneficiary_id,
                        amount_type,
                        value,
                    )

                    # This deduction should be active
                    if beneficiary_id_str in existing_map:
                        # Update existing deduction
                        existing_deduction = existing_map[beneficiary_id_str]

                        # Check if any values changed
                        old_type = existing_deduction.amount_type
                        old_value = existing_deduction.value
                        old_active = existing_deduction.is_active

                        changes_made = False
                        if old_type != amount_type:
                            current_app.logger.debug(
                                "Amount type changed: %s -> %s",
                                old_type,
                                amount_type,
                            )
                            existing_deduction.amount_type = amount_type
                            changes_made = True

                        if amount_type == 'Calculated':
                            if old_value is not None:
                                current_app.logger.debug("Clearing value for calculated deduction")
                                existing_deduction.value = None
                                changes_made = True
                        else:
                            if old_value is None or abs(float(old_value) - value) > 0.01:
                                current_app.logger.debug(
                                    "Value changed: %s -> %s", old_value, value
                                )
                                existing_deduction.value = value
                                changes_made = True

                        if not old_active:
                            current_app.logger.debug("Reactivating deduction")
                            existing_deduction.is_active = True
                            changes_made = True

                        if changes_made:
                            existing_deduction.notes = "Updated via employee edit"
                            existing_deduction.updated_at = datetime.utcnow()
                            deductions_updated += 1
                            current_app.logger.debug(
                                "Deduction updated for beneficiary %s", beneficiary_id
                            )
                        else:
                            current_app.logger.debug(
                                "No changes for beneficiary %s", beneficiary_id
                            )
                    else:
                        # Create new deduction
                        new_deduction = EmployeeRecurringDeduction()
                        new_deduction.employee_id = employee_id
                        new_deduction.beneficiary_id = int(beneficiary_id)
                        new_deduction.amount_type = amount_type
                        if amount_type == 'Calculated':
                            new_deduction.value = None
                        else:
                            new_deduction.value = value
                        new_deduction.notes = "Added via employee edit"
                        new_deduction.is_active = True
                        new_deduction.effective_date = employee.start_date

                        db.session.add(new_deduction)
                        deductions_added += 1
                        current_app.logger.debug(
                            "New deduction added for beneficiary %s", beneficiary_id
                        )
                
                # Deactivate any existing deductions that weren't in the form (removed from UI)
                for beneficiary_id_str, existing_deduction in existing_map.items():
                    if beneficiary_id_str not in processed_beneficiaries and existing_deduction.is_active:
                        existing_deduction.is_active = False
                        existing_deduction.notes = "Removed via employee edit"
                        existing_deduction.updated_at = datetime.utcnow()
                        deductions_removed += 1
                        current_app.logger.debug(
                            "Deduction removed for beneficiary %s", beneficiary_id_str
                        )
                
                current_app.logger.debug(
                    "Deduction summary - Added: %s, Updated: %s, Removed: %s",
                    deductions_added,
                    deductions_updated,
                    deductions_removed,
                )
            
            db.session.commit()
            current_app.logger.debug(
                "Recurring deductions committed for employee %s", employee_id
            )
            # Enhanced success message with deduction updates info
            success_message = f'Employee {first_name} {last_name} has been successfully updated!'
            
            # Add deduction change details if any changes were made
            deduction_changes = []
            if deductions_added > 0:
                deduction_changes.append(f'{deductions_added} deduction{"s" if deductions_added != 1 else ""} added')
            if deductions_updated > 0:
                deduction_changes.append(f'{deductions_updated} deduction{"s" if deductions_updated != 1 else ""} updated')
            if deductions_removed > 0:
                deduction_changes.append(f'{deductions_removed} deduction{"s" if deductions_removed != 1 else ""} removed')
            
            if deduction_changes:
                success_message += f' ({", ".join(deduction_changes)})'
            
            # Check if this is an AJAX request (modal submission)
            if is_ajax:
                return jsonify({
                    'success': True,
                    'message': success_message,
                    'employee_id': employee_id,
                    'redirect_url': url_for('employees.view', employee_id=employee_id)
                })
            else:
                flash(success_message, 'success')
                return redirect(url_for('employees.view', employee_id=employee_id))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error("Exception updating employee: %s", e)
            import traceback
            current_app.logger.error("Traceback: %s", traceback.format_exc())
            
            error_message = f'An error occurred while updating the employee: {str(e)}'
            
            # Check if this is an AJAX request (modal submission)
            if is_ajax:
                return jsonify({
                    'success': False,
                    'errors': [error_message],
                    'exception': str(e)
                }), 500
            else:
                flash(error_message, 'error')
                return redirect(url_for('employees.view', employee_id=employee_id))
    
    # GET request - redirect to view page (modal-based editing)
    return redirect(url_for('employees.view', employee_id=employee_id))

@employees_bp.route('/<int:employee_id>/delete', methods=['POST'])
def delete(employee_id):
    """Delete employee and all related records"""
    from app.models.ui19_record import UI19Record
    from app.models.payroll_entry import PayrollEntry
    from app.models.employee_recurring_deduction import EmployeeRecurringDeduction
    from app.models.employee_medical_aid_info import EmployeeMedicalAidInfo
    
    employee = Employee.query.get(employee_id)
    if not employee:
        flash('Employee not found.', 'error')
        return redirect(url_for('employees.index'))
    
    try:
        employee_name = employee.full_name
        
        # Delete related records first (in order to avoid foreign key constraints)
        # 1. Delete UI19 records
        UI19Record.query.filter_by(employee_id=employee_id).delete()
        
        # 2. Delete payroll entries
        PayrollEntry.query.filter_by(employee_id=employee_id).delete()
        
        # 3. Delete recurring deductions
        EmployeeRecurringDeduction.query.filter_by(employee_id=employee_id).delete()
        
        # 4. Delete medical aid info
        EmployeeMedicalAidInfo.query.filter_by(employee_id=employee_id).delete()
        
        # 5. Finally delete the employee
        db.session.delete(employee)
        db.session.commit()
        
        flash(f'Employee {employee_name} and all related records have been successfully deleted.', 'success')
        current_app.logger.info(f"Employee {employee_name} (ID: {employee_id}) deleted by user {current_user.id}")
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error("Error deleting employee: %s", e)
        flash('An error occurred while deleting the employee. Please try again or contact support.', 'error')
    
    return redirect(url_for('employees.index'))

@employees_bp.route('/department/<department_name>')
def department(department_name):
    """View employees by department"""
    
    employees = EmployeeService.get_employees_by_department(department_name)
    
    if not employees:
        flash(f'No employees found in {department_name} department.', 'info')
    
    return render_template('employees/department.html',
                         employees=employees,
                         department_name=department_name)

@employees_bp.route('/<int:employee_id>/payroll/new', methods=['GET', 'POST'])
def payroll_new(employee_id):
    """Enter new payroll data for employee"""
    
    employee = Employee.query.get(employee_id)
    if not employee:
        flash('Employee not found.', 'error')
        return redirect(url_for('employees.index'))
    
    if request.method == 'POST':
        # Get form data
        pay_period_start = request.form.get('pay_period_start', '').strip()
        pay_period_end = request.form.get('pay_period_end', '').strip()
        ordinary_hours = request.form.get('ordinary_hours', '').strip()
        overtime_hours = request.form.get('overtime_hours', '0').strip()
        sunday_hours = request.form.get('sunday_hours', '0').strip()
        public_holiday_hours = request.form.get('public_holiday_hours', '0').strip()
        hourly_rate = request.form.get('hourly_rate', '').strip()
        allowances = request.form.get('allowances', '0').strip()
        union_fee = request.form.get('union_fee', '0').strip()
        
        # Validate required fields
        errors = []
        if not pay_period_start:
            errors.append("Pay period start date is required")
        if not pay_period_end:
            errors.append("Pay period end date is required")
        if not ordinary_hours:
            errors.append("Ordinary hours worked is required")
        if not hourly_rate:
            errors.append("Hourly rate is required")
        
        # Validate dates
        pay_start_obj = None
        pay_end_obj = None
        try:
            pay_start_obj = datetime.strptime(pay_period_start, '%Y-%m-%d').date()
        except:
            errors.append("Invalid pay period start date format")
        
        try:
            pay_end_obj = datetime.strptime(pay_period_end, '%Y-%m-%d').date()
        except:
            errors.append("Invalid pay period end date format")
        
        if pay_start_obj and pay_end_obj and pay_start_obj >= pay_end_obj:
            errors.append("Pay period start date must be before end date")
        
        # Validate numeric fields
        numeric_fields = {
            'ordinary_hours': ordinary_hours,
            'overtime_hours': overtime_hours,
            'sunday_hours': sunday_hours,
            'public_holiday_hours': public_holiday_hours,
            'hourly_rate': hourly_rate,
            'allowances': allowances,
            'union_fee': union_fee
        }
        
        validated_values = {}
        for field, value in numeric_fields.items():
            try:
                decimal_value = Decimal(value) if value else Decimal('0')
                if decimal_value < 0:
                    errors.append(f"{field.replace('_', ' ').title()} cannot be negative")
                validated_values[field] = decimal_value
            except:
                errors.append(f"Invalid {field.replace('_', ' ')} value")
        
        # Check for duplicate payroll entry for the same period
        existing_entry = PayrollEntry.query.filter_by(
            employee_id=employee_id,
            pay_period_start=pay_start_obj,
            pay_period_end=pay_end_obj
        ).first()
        
        if existing_entry:
            errors.append("Payroll entry already exists for this pay period")
        
        if errors:
            for error in errors:
                flash(error, 'error')
            return render_template('payroll/new.html', employee=employee)
        
        try:
            # Handle different calculation logic for monthly vs hourly employees
            if employee.salary_type == 'monthly':
                # For monthly employees: auto-calculate hourly rate and set ordinary hours to 160
                calculated_hourly_rate = employee.salary / 160
                ordinary_hours_value = Decimal('160')  # Standard monthly hours
                hourly_rate_value = calculated_hourly_rate
            else:
                # For hourly employees: use provided values
                ordinary_hours_value = validated_values['ordinary_hours']
                hourly_rate_value = validated_values['hourly_rate']
            
            # Create new payroll entry
            payroll_entry = PayrollEntry()
            payroll_entry.employee_id = employee_id
            payroll_entry.pay_period_start = pay_start_obj
            payroll_entry.pay_period_end = pay_end_obj
            payroll_entry.month_year = datetime.now().strftime('%Y-%m')
            payroll_entry.ordinary_hours = ordinary_hours_value
            payroll_entry.overtime_hours = validated_values['overtime_hours']
            payroll_entry.sunday_hours = validated_values['sunday_hours']
            payroll_entry.public_holiday_hours = validated_values['public_holiday_hours']
            payroll_entry.hourly_rate = hourly_rate_value
            payroll_entry.allowances = validated_values['allowances']
            payroll_entry.deductions_other = Decimal('0')
            payroll_entry.union_fee = validated_values['union_fee']
            
            # Add medical aid calculations
            if employee.medical_aid_employer:
                payroll_entry.fringe_benefit_medical = employee.medical_aid_employer
            
            if employee.medical_aid_dependants:
                payroll_entry.medical_aid_tax_credit = payroll_entry.calculate_medical_tax_credit(employee.medical_aid_dependants)
            
            # Calculate statutory deductions
            payroll_entry.calculate_statutory_deductions()
            
            # Save to database
            db.session.add(payroll_entry)
            db.session.commit()
            
            flash(f'Payroll entry created successfully for {employee.full_name}. Net pay: R{payroll_entry.net_pay:,.2f}', 'success')
            return redirect(url_for('employees.view', employee_id=employee_id))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error("Error creating payroll entry: %s", e)
            flash('An error occurred while creating the payroll entry.', 'error')
            return render_template('payroll/new.html', employee=employee)
    
    # GET request - show form with default values
    # Set default pay period (current month)
    today = date.today()
    month_start = today.replace(day=1)
    if today.month == 12:
        month_end = today.replace(year=today.year + 1, month=1, day=1) - timedelta(days=1)
    else:
        month_end = today.replace(month=today.month + 1, day=1) - timedelta(days=1)
    
    default_data = {
        'pay_period_start': month_start.strftime('%Y-%m-%d'),
        'pay_period_end': month_end.strftime('%Y-%m-%d'),
        'ordinary_hours': '160',  # Standard monthly hours
        'hourly_rate': str(employee.salary) if employee.salary_type == 'hourly' else '0',
        'allowances': '500',  # Default allowance
        'overtime_hours': '0',
        'sunday_hours': '0',
        'public_holiday_hours': '0',
        'union_fee': '0'
    }
    
    return render_template('payroll/new.html', employee=employee, default_data=default_data)


@employees_bp.route('/<int:employee_id>/deductions/api')
@login_required
def employee_deductions_api(employee_id):
    """API endpoint to fetch employee recurring deductions"""
    selected_company_id = session.get('selected_company_id')
    
    if not selected_company_id:
        return jsonify({'error': 'No company selected'}), 400
    
    # Verify user has access to this company
    if not current_user.has_company_access(selected_company_id):
        return jsonify({'error': 'Access denied'}), 403
    
    # Get employee and verify it belongs to the selected company
    employee = Employee.query.filter_by(
        id=employee_id, 
        company_id=selected_company_id
    ).first()
    
    if not employee:
        return jsonify({'error': 'Employee not found'}), 404
    
    try:
        # Get employee's active recurring deductions with beneficiary information
        deductions = db.session.query(
            EmployeeRecurringDeduction, Beneficiary
        ).join(
            Beneficiary, EmployeeRecurringDeduction.beneficiary_id == Beneficiary.id
        ).filter(
            EmployeeRecurringDeduction.employee_id == employee_id,
            EmployeeRecurringDeduction.is_active == True
        ).order_by(
            Beneficiary.name.asc()
        ).all()
        
        result = []
        for deduction, beneficiary in deductions:
            result.append({
                'id': deduction.id,
                'beneficiary_id': beneficiary.id,
                'beneficiary_name': beneficiary.name,
                'beneficiary_type': beneficiary.type,
                'amount_type': deduction.amount_type.lower(),
                'value': float(deduction.value),
                'notes': deduction.notes,
                'effective_date': deduction.effective_date.isoformat() if deduction.effective_date else None
            })
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to fetch employee deductions: {str(e)}'}), 500


@employees_bp.route('/<int:employee_id>/api')
@login_required
def employee_api(employee_id):
    """API endpoint to fetch employee data for editing"""
    selected_company_id = session.get('selected_company_id')
    
    if not selected_company_id:
        return jsonify({'error': 'No company selected'}), 400
    
    # Verify user has access to this company
    if not current_user.has_company_access(selected_company_id):
        return jsonify({'error': 'Access denied'}), 403
    
    # Get employee and verify it belongs to the selected company
    employee = Employee.query.filter_by(
        id=employee_id, 
        company_id=selected_company_id
    ).first()
    
    if not employee:
        return jsonify({'error': 'Employee not found'}), 404
    
    try:
        # Convert employee data to dictionary for JSON response
        # Get employee's recurring deductions
        recurring_deductions = EmployeeRecurringDeduction.query.filter_by(
            employee_id=employee_id,
            is_active=True
        ).all()
        
        # Get all beneficiaries for this company
        beneficiaries = Beneficiary.query.filter_by(
            company_id=selected_company_id
        ).all()
        
        # Prepare recurring deductions data
        deductions_data = []
        # Pre-calculate medical aid deduction once to avoid recalculation logic
        pre_calc_medical = calculate_medical_aid_deduction(employee)
        for deduction in recurring_deductions:
            beneficiary = next((b for b in beneficiaries if b.id == deduction.beneficiary_id), None)
            if beneficiary:
                value = float(deduction.value) if deduction.value is not None else 0.0
                if deduction.amount_type == 'Calculated' and beneficiary.type == 'Medical Aid':
                    # Use pre-computed amount from medical aid info card
                    value = float(pre_calc_medical)
                deductions_data.append({
                    'id': deduction.id,
                    'beneficiary_id': deduction.beneficiary_id,
                    'beneficiary_name': beneficiary.name,
                    'deduction_type': beneficiary.type,
                    'amount_type': deduction.amount_type,
                    'value': value,
                    'enabled': deduction.is_active,
                    'is_active': deduction.is_active,
                    'notes': deduction.notes
            })

        # Calculate fringe benefit based on medical aid info
        fringe_benefit_amount = calculate_medical_aid_fringe_benefit(employee.medical_aid_info)

        # Determine medical aid tax credit according to configuration
        tax_credit_amount = 0
        info = getattr(employee, 'medical_aid_info', None)
        if info and info.use_sars_calculation:
            tax_credit_amount = calculate_medical_aid_deduction(employee)
        medical_use_sars = info.use_sars_calculation if info else False
        linked_benef_name = info.linked_beneficiary.name if info and info.linked_beneficiary else None

        # Prepare beneficiaries data for dropdowns
        beneficiaries_data = []
        for beneficiary in beneficiaries:
            beneficiaries_data.append({
                'id': beneficiary.id,
                'name': beneficiary.name,
                'type': beneficiary.type
            })

        employee_data = {
            'id': employee.id,
            'employee_id': employee.employee_id,
            'first_name': employee.first_name,
            'last_name': employee.last_name,
            'id_number': employee.id_number,
            'identification_type': employee.identification_type,
            'passport_number': employee.passport_number,
            'date_of_birth': employee.date_of_birth.isoformat() if employee.date_of_birth else '',
            'gender': employee.gender,
            'marital_status': employee.marital_status,
            'tax_number': employee.tax_number,
            'cell_number': employee.cell_number,
            'email': employee.email,
            'physical_address': employee.physical_address,
            
            # Employment Information
            'department': employee.department,
            'job_title': employee.job_title,
            'start_date': employee.start_date.isoformat() if employee.start_date else '',
            'end_date': employee.end_date.isoformat() if employee.end_date else '',
            'employment_type': employee.employment_type,
            'reporting_manager': employee.reporting_manager,
            'union_member': employee.union_member,
            'union_name': employee.union_name,
            
            # Compensation
            'salary_type': employee.salary_type,
            'salary': float(employee.salary) if employee.salary else 0,
            'overtime_eligible': employee.overtime_eligible,
            'ordinary_hours_per_day': float(employee.ordinary_hours_per_day),
            'work_days_per_month': employee.work_days_per_month,
            'overtime_calc_method': employee.overtime_calc_method,
            'allowances': float(employee.allowances) if employee.allowances else 0,
            'bonus_type': employee.bonus_type,
            
            # Statutory Deductions
            'uif_contributing': employee.uif_contributing,
            'sdl_contributing': employee.sdl_contributing,
            'paye_exempt': employee.paye_exempt,
            
            # Medical Aid
            'medical_aid_member': employee.medical_aid_member,

            'medical_aid_fringe_benefit': float(employee.medical_aid_fringe_benefit) if employee.medical_aid_fringe_benefit else 0,
            'fringe_benefit_amount': fringe_benefit_amount,
            'medical_aid_tax_credit': tax_credit_amount,
            'medical_aid_use_sars': medical_use_sars,
            'medical_aid_linked_beneficiary': linked_benef_name,
            
            # Banking
            'bank_name': employee.bank_name,
            'account_number': employee.account_number,
            'account_type': employee.account_type,
            'annual_leave_days': employee.annual_leave_days or 15,
            
            # Recurring Deductions
            'recurring_deductions': deductions_data,
            'beneficiaries': beneficiaries_data
        }
        
        return jsonify(employee_data)
        
    except Exception as e:
        return jsonify({'error': f'Failed to fetch employee data: {str(e)}'}), 500


@employees_bp.route('/api/employees/<int:employee_id>/payroll-ytd', methods=['GET'])
@login_required
def get_employee_ytd(employee_id):
    """Return year-to-date payroll totals for an employee."""

    # Ensure employee belongs to the currently selected company
    employee = Employee.query.get(employee_id)
    selected_company_id = session.get('selected_company_id')

    if not employee or employee.company_id != selected_company_id:
        return jsonify({'error': 'Unauthorized'}), 403

    # Determine start of the current SARS tax year (dynamic from SARS config)
    from app.services.sars_service import SARSService
    sars_config = SARSService.get_global_sars_config()
    # Tax year starts: sars_config.tax_year_start_display
    today = date.today()
    tax_year_start = date(today.year if today.month >= 3 else today.year - 1, 3, 1)

    # Calculate the YTD totals
    totals = calculate_ytd_totals(employee_id, tax_year_start)
    if not totals:
        return jsonify({'error': 'Employee not found'}), 404

    return jsonify(totals)


@employees_bp.route('/update_medical_aid/<int:employee_id>', methods=['POST'])
@login_required
def update_medical_aid(employee_id):
    """Update an employee's medical aid information."""

    selected_company_id = session.get('selected_company_id')
    employee = Employee.query.get(employee_id)

    if not employee:
        abort(404)

    if selected_company_id and employee.company_id != selected_company_id:
        abort(403)

    employee.medical_aid_scheme = request.form.get('medical_aid_scheme', '').strip() or None
    employee.medical_aid_number = request.form.get('medical_aid_number', '').strip() or None
    employee.medical_aid_dependants = int(request.form.get('medical_aid_dependants', '0') or 0)
    employee.medical_aid_principal_member = request.form.get('medical_aid_principal_member') == 'on'

    linked_id = request.form.get('linked_medical_beneficiary_id')
    if linked_id:
        employee.linked_medical_beneficiary_id = int(linked_id)
    elif not employee.linked_medical_beneficiary_id:
        deduction = (
            EmployeeRecurringDeduction.query
            .join(Beneficiary)
            .filter(
                EmployeeRecurringDeduction.employee_id == employee.id,
                EmployeeRecurringDeduction.is_active.is_(True),
                Beneficiary.type == 'Medical Aid',
            )
            .order_by(EmployeeRecurringDeduction.id)
            .first()
        )
        if deduction:
            employee.linked_medical_beneficiary_id = deduction.beneficiary_id

    info = employee.medical_aid_info
    if not info:
        info = EmployeeMedicalAidInfo(employee_id=employee.id)
        db.session.add(info)

    info.scheme_name = employee.medical_aid_scheme
    info.membership_number = employee.medical_aid_number
    info.number_of_dependants = employee.medical_aid_dependants
    info.main_member = employee.medical_aid_principal_member
    info.linked_beneficiary_id = employee.linked_medical_beneficiary_id
    info.additional_dependants = int(request.form.get('additional_dependants', '0') or 0)
    info.employer_contribution_override = request.form.get('employer_contribution_override') or None
    info.employee_contribution_override = request.form.get('employee_contribution_override') or None
    info.use_sars_calculation = request.form.get('use_sars_calculation') == 'on'
    eff = request.form.get('effective_from')
    if eff:
        info.effective_from = datetime.strptime(eff, '%Y-%m-%d').date()
    else:
        info.effective_from = None

    db.session.commit()
    flash('Medical aid information updated.', 'success')
    return redirect(url_for('employees.view', employee_id=employee_id))

@employees_bp.route('/import', methods=['GET', 'POST'])
@login_required
def import_employees():
    """Upload file and preview employees before import"""
    if request.method == 'POST':
        file = request.files.get('file')
        apply_defaults = bool(request.form.get('apply_defaults'))
        if not file:
            flash('No file uploaded', 'error')
            return redirect(url_for('employees.import_employees'))
        try:
            df = pd.read_excel(file, dtype=str)
        except Exception as e:
            flash(f'Unable to read file: {e}', 'error')
            return redirect(url_for('employees.import_employees'))

        rows = []
        for index, raw_row in df.iterrows():
            data = {}
            errors = []
            for hdr, field in IMPORT_HEADER_MAP.items():
                raw_val = raw_row.get(hdr, '')
                data[field] = clean_import_value(raw_val)
            
            # Validate and normalize cell number
            cell_number = data.get('cell_number', '').strip()
            if cell_number:
                cell_valid, cell_error = validate_sa_cell_number(cell_number)
                if not cell_valid:
                    errors.append(f"Cell number: {cell_error}")
                else:
                    data['cell_number'] = normalize_phone(cell_number)
            else:
                errors.append("Cell number is required")
            if data.get('gender'):
                g = data['gender'].strip().lower()
                data['gender'] = GENDER_MAP.get(g, g.title())
            if data.get('marital_status'):
                m = data['marital_status'].strip().lower()
                data['marital_status'] = STATUS_MAP.get(m, m.title())
            if data.get('bank_name'):
                b = data['bank_name'].strip().lower()
                for k, v in BANK_MAP.items():
                    if k in b:
                        data['bank_name'] = v
                        break
            id_number = data.get('id_number')
            if id_number:
                valid, msg = validate_south_african_id(str(id_number))
                if not valid:
                    errors.append(msg)
                else:
                    birth, gen = extract_info_from_id(str(id_number))
                    if birth and not data.get('date_of_birth'):
                        data['date_of_birth'] = birth
                    if gen and not data.get('gender'):
                        data['gender'] = gen
            if data.get('date_of_birth'):
                dob = parse_date(data['date_of_birth'])
                if dob:
                    data['date_of_birth'] = dob
                else:
                    errors.append('Invalid date of birth')
                    data['date_of_birth'] = None
            if data.get('start_date'):
                sd = parse_date(data['start_date'])
                if sd:
                    data['start_date'] = sd
                else:
                    errors.append('Invalid start date')
                    data['start_date'] = None
            rows.append({'data': data, 'errors': errors, 'status': 'pending'})
        session['import_rows'] = rows
        session['apply_defaults'] = apply_defaults
        session['field_order'] = list(IMPORT_HEADER_MAP.values())
        return redirect(url_for('employees.import_review'))
    return render_template('employees/import.html')

@employees_bp.route('/import/review', methods=['GET'])
@login_required
def import_review():
    """Review parsed employees before final import"""
    rows = session.get('import_rows')
    if not rows:
        flash('No import data found. Please upload a file.', 'warning')
        return redirect(url_for('employees.import_employees'))
    field_order = session.get('field_order', list(IMPORT_HEADER_MAP.values()))
    return render_template('employees/import_review.html', rows=rows, field_order=field_order, apply_defaults=session.get('apply_defaults'))

@employees_bp.route('/import/confirm', methods=['POST'])
@login_required
def import_confirm():
    """Import employees after preview"""
    rows = session.get('import_rows', [])
    apply_defaults = session.get('apply_defaults', False)
    selected_company_id = session.get('selected_company_id')
    company = Company.query.get(selected_company_id) if selected_company_id else None
    imported = 0
    skipped = 0
    error_count = 0
    for idx, row in enumerate(rows):
        status = request.form.get(f'status_{idx}', row.get('status'))
        if status != 'confirmed':
            skipped += 1
            continue
        data = {}
        for field in IMPORT_HEADER_MAP.values():
            form_key = f'rows-{idx}-{field}'
            raw_val = request.form.get(form_key) or row['data'].get(field)
            data[field] = clean_import_value(raw_val)
        data['cell_number'] = normalize_phone(data.get('cell_number'))
        if apply_defaults and company:
            default_map = {
                'salary_type': 'default_salary_type',
                'monthly_salary': 'default_salary',
                'annual_leave_days': 'default_annual_leave_days',
                'overtime_multiplier': 'overtime_multiplier',
                'sunday_multiplier': 'sunday_multiplier',
                'holiday_multiplier': 'public_holiday_multiplier',
            }
            for f, attr in default_map.items():
                if not data.get(f) and getattr(company, attr, None) is not None:
                    data[f] = getattr(company, attr)
        try:
            employee = Employee()
            employee.company_id = selected_company_id
            employee.employee_id = generate_employee_id(selected_company_id)
            employee.first_name = data.get('first_name')
            employee.last_name = data.get('last_name')
            employee.id_number = str(data.get('id_number') or '')
            employee.tax_number = str(data.get('tax_number')) if data.get('tax_number') else None
            employee.cell_number = data.get('cell_number')
            employee.email = data.get('email') or None
            employee.date_of_birth = parse_date(data.get('date_of_birth')) if data.get('date_of_birth') else None
            employee.gender = data.get('gender') or None
            employee.marital_status = data.get('marital_status') or None
            employee.department = data.get('department') or ''
            employee.job_title = data.get('job_title') or ''
            employee.start_date = parse_date(data.get('start_date')) if data.get('start_date') else None
            employee.employment_type = data.get('employment_type') or 'Full-Time'
            stype = (data.get('salary_type') or 'monthly').lower()
            employee.salary_type = stype
            salary_val = data.get('monthly_salary') if stype == 'monthly' else data.get('hourly_rate')
            employee.salary = Decimal(str(salary_val)) if salary_val else Decimal('0')
            employee.bank_name = data.get('bank_name') or ''
            employee.account_number = str(data.get('account_number') or '')
            employee.account_type = data.get('account_type') or 'Savings'
            al = data.get('annual_leave_days') or 15
            employee.annual_leave_days = int(al)
            db.session.add(employee)
            db.session.commit()
            imported += 1
        except Exception as e:
            db.session.rollback()
            error_count += 1
            flash(f'Failed to import row {idx + 1}: {e}', 'error')
    flash(f'Import complete: {imported} imported, {skipped} skipped, {error_count} errors', 'success')
    return redirect(url_for('employees.index'))


@employees_bp.route('/<int:employee_id>/terminate', methods=['POST'])
@login_required
def terminate(employee_id):
    """Terminate an employee and generate UI19 record"""
    # Get current company
    current_company_id = session.get('current_company_id')
    if not current_company_id:
        flash('Please select a company first.', 'error')
        return redirect(url_for('dashboard.overview'))
    
    # Get employee with company access control
    employee = Employee.query.filter_by(id=employee_id, company_id=current_company_id).first_or_404()
    
    # Check if employee is already terminated
    if employee.employment_status == 'Terminated':
        flash(f'{employee.full_name} is already terminated.', 'warning')
        return redirect(url_for('employees.view', employee_id=employee_id))
    
    try:
        # Get form data
        termination_reason = request.form.get('termination_reason')
        termination_date_str = request.form.get('termination_date')
        termination_notes = request.form.get('termination_notes', '').strip()
        
        # Validate required fields
        if not termination_reason:
            flash('Termination reason is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        if not termination_date_str:
            flash('Termination date is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Parse termination date
        try:
            termination_date = datetime.strptime(termination_date_str, '%Y-%m-%d').date()
        except ValueError:
            flash('Invalid termination date format.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Validate termination date is not before start date
        if termination_date < employee.start_date:
            flash('Termination date cannot be before the employee start date.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Update employee record
        employee.employment_status = 'Terminated'
        employee.end_date = termination_date
        employee.termination_reason = termination_reason
        
        # Create UI19 record
        ui19_record = UI19Record(
            employee_id=employee.id,
            company_id=current_company_id,
            start_date=employee.start_date,
            end_date=termination_date,
            termination_reason=termination_reason,
            notes=termination_notes if termination_notes else None,
            status='Generated',
            created_by=current_user.id
        )
        
        # Save to database
        db.session.add(ui19_record)
        db.session.commit()
        
        flash(f'{employee.full_name} has been successfully terminated. UI19 record generated.', 'success')
        return redirect(url_for('employees.view', employee_id=employee_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error terminating employee {employee_id}: {str(e)}')
        flash(f'Error terminating employee: {str(e)}', 'error')
        return redirect(url_for('employees.view', employee_id=employee_id))


@employees_bp.route('/<int:employee_id>/reinstate', methods=['POST'])
@login_required
def reinstate(employee_id):
    """Reinstate a terminated employee"""
    # Get current company
    current_company_id = session.get('current_company_id')
    if not current_company_id:
        flash('Please select a company first.', 'error')
        return redirect(url_for('dashboard.overview'))
    
    # Get employee with company access control
    employee = Employee.query.filter_by(id=employee_id, company_id=current_company_id).first_or_404()
    
    # Check if employee is terminated
    if employee.employment_status != 'Terminated':
        flash(f'{employee.full_name} is not terminated and cannot be reinstated.', 'warning')
        return redirect(url_for('employees.view', employee_id=employee_id))
    
    try:
        # Get form data
        start_date_str = request.form.get('start_date')
        employment_type = request.form.get('employment_type')
        department = request.form.get('department')
        job_title = request.form.get('job_title')
        reinstatement_notes = request.form.get('reinstatement_notes', '').strip()
        
        # Validate required fields
        if not start_date_str:
            flash('Start date is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        if not employment_type:
            flash('Employment type is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
            
        if not department:
            flash('Department is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
            
        if not job_title:
            flash('Job title is required.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Parse start date
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        except ValueError:
            flash('Invalid start date format.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Update employee record
        employee.employment_status = 'Active'
        employee.start_date = start_date
        employee.employment_type = employment_type
        employee.department = department
        employee.job_title = job_title
        employee.end_date = None
        employee.termination_reason = None
        
        # Save to database
        db.session.commit()
        
        flash(f'{employee.full_name} has been successfully reinstated to active status.', 'success')
        return redirect(url_for('employees.view', employee_id=employee_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error reinstating employee {employee_id}: {str(e)}')
        flash(f'Error reinstating employee: {str(e)}', 'error')
        return redirect(url_for('employees.view', employee_id=employee_id))


@employees_bp.route('/<int:employee_id>/generate-uif-doc')
@login_required
def generate_uif_doc(employee_id):
    """Generate UIF Declaration PDF document for terminated employee"""
    # Get current company
    current_company_id = session.get('current_company_id')
    if not current_company_id:
        flash('Please select a company first.', 'error')
        return redirect(url_for('dashboard.overview'))
    
    # Get employee with company access control
    employee = Employee.query.filter_by(id=employee_id, company_id=current_company_id).first_or_404()
    company = employee.company
    
    # Check if employee has termination records
    ui19_records = UI19Record.query.filter_by(employee_id=employee_id, company_id=current_company_id).all()
    if not ui19_records:
        flash('No UI19 termination records found for this employee.', 'error')
        return redirect(url_for('employees.view', employee_id=employee_id))
    
    try:
        # Get the UI19 template from database
        ui19_template = DocumentTemplate.query.filter(
            DocumentTemplate.document_type.ilike('%UI19%')
        ).first()
        if not ui19_template:
            flash('UI19 document template not found. Please contact system administrator.', 'error')
            return redirect(url_for('employees.view', employee_id=employee_id))
        
        # Create temporary file for template
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_template:
            temp_template.write(ui19_template.file_data)
            template_path = temp_template.name
        
        # Load the template document
        doc = Document(template_path)
        
        # Helper function to replace text in document
        def fill_field(doc, placeholder, value):
            """Replace placeholder text in document paragraphs and tables"""
            value = str(value) if value is not None else ''
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if placeholder in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)
        
        # Get latest termination record
        latest_record = ui19_records[-1]  # Most recent record
        
        # Prepare company data
        uif_ref = company.uif_reference_number or ''
        paye_ref = company.paye_reference_number or ''
        reg_number = company.registration_number or ''
        contact_email = company.email or ''
        phone_number = company.phone or ''
        emp_fname = company.employer_first_name or ''
        emp_lname = company.employer_last_name or ''
        emp_id = company.employer_id_number or ''
        
        # Prepare employee data
        emp_last = employee.last_name or ''
        emp_initials = employee.first_name[0] if employee.first_name else ''
        emp_id_no = employee.id_number or ''
        
        # Get latest payroll entry for salary information
        latest_payroll = PayrollEntry.query.filter_by(employee_id=employee_id).order_by(PayrollEntry.pay_period_start.desc()).first()
        if latest_payroll:
            emp_salary = str(int(float(latest_payroll.gross_pay or 0)))
            emp_salary_cents = "00"
        else:
            emp_salary = str(int(float(employee.salary or 0)))
            emp_salary_cents = "00"
        
        # Calculate average monthly hours (default to 160)
        emp_hours = "160"  # Standard full-time hours per month
        
        # Format dates
        start_date = employee.start_date.strftime('%d%m%y') if employee.start_date else ''
        end_date = latest_record.end_date.strftime('%d%m%y') if latest_record.end_date else ''
        
        # Termination details
        term_reason = "1"  # Default termination code - can be enhanced based on termination_reason
        uif_status = "Y"   # Assume UIF eligible unless specified
        uif_reason = "1"   # Default UIF reason code
        
        # Current date info
        today = datetime.today().strftime('%d/%m/%Y')
        month_str = datetime.today().strftime('%B').upper()
        
        # Field mapping for UI19 form
        field_map = {
            # UIF Reference Number (split into individual characters)
            'E_1': uif_ref[0:1] if len(uif_ref) > 0 else '',
            'E_2': uif_ref[1:2] if len(uif_ref) > 1 else '',
            'E_3': uif_ref[2:3] if len(uif_ref) > 2 else '',
            'E_4': uif_ref[3:4] if len(uif_ref) > 3 else '',
            'E_5': uif_ref[4:5] if len(uif_ref) > 4 else '',
            'E_6': uif_ref[5:6] if len(uif_ref) > 5 else '',
            'E_7': uif_ref[6:7] if len(uif_ref) > 6 else '',
            'E_8': uif_ref[7:8] if len(uif_ref) > 7 else '',
            'E_9': uif_ref[9:10] if len(uif_ref) > 9 else '',
            
            # Period information
            'ED_MONTH': month_str,
            
            # PAYE Reference Number (split into individual characters)
            'P_1': paye_ref[0:1] if len(paye_ref) > 0 else '',
            'P_2': paye_ref[1:2] if len(paye_ref) > 1 else '',
            'P_3': paye_ref[2:3] if len(paye_ref) > 2 else '',
            'P_4': paye_ref[3:4] if len(paye_ref) > 3 else '',
            'P_5': paye_ref[4:5] if len(paye_ref) > 4 else '',
            'P_6': paye_ref[5:6] if len(paye_ref) > 5 else '',
            'P_7': paye_ref[6:7] if len(paye_ref) > 6 else '',
            'P_8': paye_ref[7:8] if len(paye_ref) > 7 else '',
            'P_9': paye_ref[8:9] if len(paye_ref) > 8 else '',
            'P_10': paye_ref[9:10] if len(paye_ref) > 9 else '',
            
            # Company Registration (split into individual characters)
            'C_1': reg_number[0:1] if len(reg_number) > 0 else '',
            'C_2': reg_number[1:2] if len(reg_number) > 1 else '',
            'C_3': reg_number[2:3] if len(reg_number) > 2 else '',
            'C_4': reg_number[3:4] if len(reg_number) > 3 else '',
            'C_5': reg_number[4:5] if len(reg_number) > 4 else '',
            'C_6': reg_number[5:6] if len(reg_number) > 5 else '',
            'C_7': reg_number[6:7] if len(reg_number) > 6 else '',
            'C_8': reg_number[7:8] if len(reg_number) > 7 else '',
            'C_9': reg_number[8:9] if len(reg_number) > 8 else '',
            'C_10': reg_number[9:10] if len(reg_number) > 9 else '',
            'C_11': reg_number[10:11] if len(reg_number) > 10 else '',
            'C_12': reg_number[11:12] if len(reg_number) > 11 else '',
            'C_13': reg_number[12:13] if len(reg_number) > 12 else '',
            'C_14': reg_number[13:14] if len(reg_number) > 13 else '',
            'C_15': reg_number[14:15] if len(reg_number) > 14 else '',
            
            # Company contact information
            'COMPANY_EMAIL': contact_email,
            'COMPANY_PHONE': phone_number,
            
            # Employer information
            'First_Last_Name': f"{emp_fname} {emp_lname}".strip(),
            
            # Employee information
            'EMP_L_NAME': emp_last,
            'EMP_I': emp_initials,
            
            # Employee ID Number (split into individual characters)
            'I_1': emp_id_no[0:1] if len(emp_id_no) > 0 else '',
            'I_2': emp_id_no[1:2] if len(emp_id_no) > 1 else '',
            'I_3': emp_id_no[2:3] if len(emp_id_no) > 2 else '',
            'I_4': emp_id_no[3:4] if len(emp_id_no) > 3 else '',
            'I_5': emp_id_no[4:5] if len(emp_id_no) > 4 else '',
            'I_6': emp_id_no[5:6] if len(emp_id_no) > 5 else '',
            'I_7': emp_id_no[6:7] if len(emp_id_no) > 6 else '',
            'I_8': emp_id_no[7:8] if len(emp_id_no) > 7 else '',
            'I_9': emp_id_no[8:9] if len(emp_id_no) > 8 else '',
            'I_10': emp_id_no[9:10] if len(emp_id_no) > 9 else '',
            'I_11': emp_id_no[10:11] if len(emp_id_no) > 10 else '',
            'I_12': emp_id_no[11:12] if len(emp_id_no) > 11 else '',
            'I_13': emp_id_no[12:13] if len(emp_id_no) > 12 else '',
            
            # Employee salary and hours
            'EMP_R': emp_salary,
            'EMP_C': emp_salary_cents,
            'EMP_H': emp_hours,
            
            # Employment dates (split into individual characters)
            'J_1': start_date[0:1] if len(start_date) > 0 else '',
            'J_2': start_date[1:2] if len(start_date) > 1 else '',
            'J_3': start_date[2:3] if len(start_date) > 2 else '',
            'J_4': start_date[3:4] if len(start_date) > 3 else '',
            'J_5': start_date[4:5] if len(start_date) > 4 else '',
            'J_6': start_date[5:6] if len(start_date) > 5 else '',
            
            'J_7': end_date[0:1] if len(end_date) > 0 else '',
            'J_8': end_date[1:2] if len(end_date) > 1 else '',
            'J_9': end_date[2:3] if len(end_date) > 2 else '',
            'J_10': end_date[3:4] if len(end_date) > 3 else '',
            'J_11': end_date[4:5] if len(end_date) > 4 else '',
            'J_12': end_date[5:6] if len(end_date) > 5 else '',
            
            # Termination codes
            'J_CODE': term_reason,
            'J_YN': uif_status,
            'J_R': uif_reason,
            
            # Signature information
            'EMP_ID': emp_id,
            'E_SIG': f"{emp_fname} {emp_lname}".strip(),
            'E_DATE': today
        }
        
        # Replace all placeholders in the document
        for placeholder, value in field_map.items():
            fill_field(doc, placeholder, value)
        
        # Create temporary files for output
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            temp_docx_path = temp_docx.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        # Save the filled document
        doc.save(temp_docx_path)
        
        # Convert DOCX to PDF
        try:
            convert(temp_docx_path, temp_pdf_path)
        except Exception as conv_error:
            current_app.logger.error(f'Error converting DOCX to PDF: {str(conv_error)}')
            # If PDF conversion fails, return the DOCX file
            filename = f'UI19_Declaration_{employee.last_name}_{employee.first_name}.docx'
            response = send_file(temp_docx_path, as_attachment=True, download_name=filename)
            
            # Clean up temporary files
            try:
                os.unlink(template_path)
                os.unlink(temp_docx_path)
            except OSError:
                pass
            
            return response
        
        # Generate filename
        filename = f'UI19_Declaration_{employee.last_name}_{employee.first_name}.pdf'
        
        # Send the PDF file
        response = send_file(temp_pdf_path, as_attachment=True, download_name=filename)
        
        # Clean up temporary files
        try:
            os.unlink(template_path)
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
        except OSError:
            pass
        
        return response
        
    except Exception as e:
        current_app.logger.error(f'Error generating UIF declaration for employee {employee_id}: {str(e)}')
        flash('Failed to generate UIF declaration document. Please try again.', 'error')
        return redirect(url_for('employees.view', employee_id=employee_id))
